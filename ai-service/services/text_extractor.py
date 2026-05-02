"""
Text extraction service for PDF and DOCX resume files.
"""
import io
from PyPDF2 import PdfReader
from docx import Document


def extract_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text content from a PDF file.
    
    Args:
        file_bytes: Raw bytes of the PDF file.
        
    Returns:
        Extracted text as a single string.
    """
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n".join(text_parts).strip()
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")


def extract_from_docx(file_bytes: bytes) -> str:
    """
    Extract text content from a DOCX file.
    
    Args:
        file_bytes: Raw bytes of the DOCX file.
        
    Returns:
        Extracted text as a single string.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n".join(text_parts).strip()
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")


def extract_text(file_bytes: bytes, file_type: str) -> str:
    """
    Extract text from a file based on its type.
    
    Args:
        file_bytes: Raw bytes of the file.
        file_type: File type ('pdf' or 'docx').
        
    Returns:
        Extracted text as a string.
    """
    file_type = file_type.lower().strip('.')
    
    if file_type == "pdf":
        return extract_from_pdf(file_bytes)
    elif file_type in ("docx", "doc"):
        return extract_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {file_type}. Supported types: pdf, docx")
